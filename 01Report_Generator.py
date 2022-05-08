from docx import Document
from docx.shared import Inches
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
#from win32com import client
################################################################################
# Definition part
df = pd.read_excel(r"C:\Users****,sheet_name=None)
Gruppo=df['COP'].loc[6][0]
df1=df['1.DB Società']
cols1=df1.loc[1]
cols1_1=cols1[1:4]
cols1_3=cols1[10:15]
cols=df1.loc[2]
cols1_2=cols[4:10]
COLS=[]
COLS[0:3]=cols1_1
COLS[3:9]=cols1_2
COLS[10:15]=cols1_3
COLS=['n']+COLS
df1.columns=COLS
df1=df1[3:]
df2=df['Criteri']
df2_1=df2.copy()
df2_1.columns=df2_1.loc[1]
df2_1=df2_1[2:7]
df2_1.index=range(0,len(df2_1))
df2_2=df2.copy()
df2_2.columns=df2_2.loc[9]
df2_2=df2_2[10:15]
df2_2.index=range(0,len(df2_2))
df3=df['2.DB ubicazioni']
df3.columns=df3.iloc[3]
df3=df3[4:]
df3=df3[['società','indirizzo','n°','frazione/località/ specifiche','cap','città','pv','nazione','link mappa/ latitudine e longitudine','Anno costruzione','bene culturale (SI/NO)','soggetto proprietario','superficie di sviluppo totale (mq) (1)','superficie di sviluppo specifica (mq) (1)','destinazione d\'uso generale (2)','destinazione d\'uso specifica (3)','N. Dipendenti','Ricavi infragruppo','Autorizzazioni ambientali AIA o AUA','Fabbricato','Contenuto (5)','Merci\n(max giacenza)','Totale','Polizza']]
df3.dropna(subset=['società'],inplace=True)
df3.index=range(0,len(df3))
df3['ID']=range(1,len(df3)+1)
cols3=['ID','società','indirizzo','n°','frazione/località/ specifiche','cap','città','pv','nazione','link mappa/ latitudine e longitudine','Anno costruzione','bene culturale (SI/NO)','soggetto proprietario','superficie di sviluppo totale (mq) (1)','superficie di sviluppo specifica (mq) (1)','destinazione d\'uso generale (2)','destinazione d\'uso specifica (3)','N. Dipendenti','Ricavi infragruppo','Autorizzazioni ambientali AIA o AUA','Fabbricato','Contenuto (5)','Merci\n(max giacenza)','Totale','Polizza']
df3=df3[cols3]
################################################################################
# Preparation Part
WW = RGBColor(255, 255, 255)
document = Document('temp.docx')

style = document.styles['Normal']
font = style.font
font.name='Arial'
#header = document.sections[0].header
#paragraph = header.paragraphs[0]
#logo_run = paragraph.add_run()
#logo_run.add_picture("1_c_logo.png", width=Inches(1.75))
#text_run = paragraph.add_run()
#paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
#text_run.text = '\t' + '\t' + "© Copyright 2019 - Strategica Risk Consulting S.r.l. - Milano" # For center align of text
text_run.style = "Body Text Char"
document.add_page_break()
################################################################################
# Adding part:
document.add_heading('PREMESSA', 1)
document.add_paragraph('Riservatezza')
T='Strategica Risk Consulting S.r.l., di seguito “Strategica”, e {} si impegnano a mantenere strettamente confidenziali e a non divulgare a terzi, senza il preventivo assenso, tutte le informazioni reciprocamente fornite o apprese durante lo sviluppo del progetto, quali, a puro titolo esemplificativo e non limitativo, materiali, documenti, studi, analisi, know-how verbale o scritto, analisi organizzative o commerciali, conversazioni, espressioni d’opinioni o descrizioni di eventi.'
p=document.add_paragraph(T.format(Gruppo))
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph('Limitazioni di responsabilità')
T='Fermo restando l\'obbligo di Strategica di fornire servizi eseguiti a regola d\'arte, le analisi e le valutazioni oggetto del presente elaborato sono basate esclusivamente su dati, valori, documenti ed informazioni (di seguito congiuntamente indicate come “Informazioni”) forniti da {} pertanto nessuna responsabilità potrà essere riconducibile a Strategica in ordine alla veridicità delle Informazioni utilizzate e ai conseguenti risultati. In particolare, {} sarà il solo responsabile delle decisioni prese in relazione alle attività intraprese con l\'assistenza e/o la consulenza del personale di Strategica e dei risultati di tali attività.'
p=document.add_paragraph(T.format(Gruppo,Gruppo))
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_page_break()
## Continue:
document.add_heading('PERIMETRO DELL\'ANALISI', 1)
if len(df1['Denominazione societaria '])==1:
    p=document.add_paragraph('L’Analisi tratta la società')
else:
    p=document.add_paragraph('L’Analisi tratta le società')

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for i in df1['Denominazione societaria ']:
    document.add_paragraph(i, style='List Bullet')

p=document.add_paragraph('Inoltre, l\'analisi riguarderà solo i rischi operativi.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_heading('SCOPO E METODOLOGIA', 1)
T='Scopo di questa analisi è l’individuazione di alcuni principali elementi di Risk Assesment al fine di supportare {} nel raggiungimento dei propri obiettivi minimizzando gli effetti del rischio al minimo costo. Questa analisi applica le linee guida contenute nella “ISO 31.000”, che prevede le seguenti fasi'
p=document.add_paragraph(T.format(Gruppo))
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph('Framework', style='List Bullet')
document.add_paragraph('Processo', style='List Bullet')
document.add_paragraph('Contesto, criteri e scopo', style='List Bullet')
document.add_paragraph('Identificazione', style='List Bullet 2')
document.add_paragraph('Analisi', style='List Bullet 2')
document.add_paragraph('Valutazione e', style='List Bullet 2')
document.add_paragraph('Trattamento dei rischi', style='List Bullet 2')
I=document.add_picture('2_c_metodology.png', width=Inches(3))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
document.add_page_break()
document.add_heading('Framework', level=2)
p=document.add_paragraph('Nel framework si definiscono gli aspetti organizzativi del processo di risk management.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_heading('Contesto', level=2)
p=document.add_paragraph('Il processo di analisi dei rischi inizia con un’attentata analisi del contesto in cui il Gruppo opera, il contesto è rappresentato dall’ambiente esterno ed interno all’azienda.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_heading('Criteri', level=2)
p=document.add_paragraph('A valle dell’analisi del contesto sono stati individuati i principali criteri, che sono')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph('Tassonomia dei rischi', style='List Bullet')
document.add_paragraph('Ponderazione dei rischi', style='List Bullet')
document.add_heading('Valutazione e trattamento', level=2)
p=document.add_paragraph('Seguono le fasi di valutazione e trattamento del rischio.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p=document.add_paragraph('La fase di valutazione identifica i rischi più significativi per il gruppo e assegna una priorità nel trattamento. Il trattamento del rischio analizzato in questo documento è il trattamento che precede il trasferimento ai mercati assicurativi, analisi fatta successivamente.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_heading('Criteri', 1)
p=document.add_paragraph('Sulla base del contesto in cui opera l’azienda sono stati individuati i seguenti criteri per classificazione dei rischi (Tassonomia) e loro ponderazione.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

R = (
    ('Danni ai beni', '-'),
    ('Responsabilità', '-'),
    ('Business Interruption e Loss of Profit', '-'),
    ('Cyber-informatico', '-'),
    ('Altri', '-')
)


table = document.add_table(rows=1, cols=2)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Categoria delle cause dei rischi'
hdr_cells[1].text = 'Sotto Categoria'

for A, B in R:
    row_cells = table.add_row().cells
    row_cells[0].text = A
    row_cells[1].text = B

table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(.3)

p=document.add_paragraph('I criteri di ponderazione utilizzati tengono presente il patrimonio del gruppo, la significatività dei rischi a cui è esposto e la durata del piano industriale.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



table = document.add_table(rows=6, cols=4)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Classe Impatto'
hdr_cells[1].text = 'Livello'
hdr_cells[2].text = 'Quantitativo'
hdr_cells[3].text = 'Qualitativo'
df2_2.columns=['nan', 'LL', 'Livello', 'Quantitativo', 'Qualitativo', 'nan', 'nan']
df2_2=df2_2.fillna('-')
for i in range(0,len(df2_2)):
    table.cell(i+1, 0).text=str(i+1)
    table.cell(i+1, 1).text=str(df2_2['Livello'].values[i])
    table.cell(i+1, 2).text=str(df2_2['Quantitativo'].values[i])
    table.cell(i+1, 3).text=str(df2_2['Qualitativo'].values[i])

table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)

p=document.add_paragraph('La stima dell’impatto si basa su due criteri congiunti (il superamento di un parametro prevale sull’altro)')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_paragraph('la gravità del danno recato al patrimonio del Gruppo', style='List Bullet')
document.add_paragraph('la gravità del danno alla reputazione', style='List Bullet')

p=document.add_paragraph('Un evento si può verificare nell’arco di tempo definito')

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
table = document.add_table(rows=6, cols=3)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Classe Probabilità'
hdr_cells[1].text = 'Livello'
hdr_cells[2].text = 'Descrizione'
df2_1.columns=['nan', 'LL', 'Livello', 'Descrizione', 'nan', 'nan', 'nan']
for i in range(0,len(df2_1)):
    table.cell(i+1, 0).text=str(i+1)
    table.cell(i+1, 1).text=str(df2_1['Livello'].values[i])
    table.cell(i+1, 2).text=str(df2_1['Descrizione'].values[i])

table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)

document.add_heading('CONTESTO', 1)
document.add_heading('Il gruppo', level=2)
# # # # This part should be completed.
if len(df1)<2:
    T='Il {} (Fig 1 e Tab. 1) è composto da una società, possedute da {}. La società opera nel {}.'
    p=document.add_paragraph(T.format(Gruppo,df1.iloc[0][1],df1['Area di business'].iloc[0]))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#else: # # # # This Part should be completed with an example
#    T='Il {} (Fig 1 e Tab. 1) è composto da diverse società, possedute da {}. Tutte le società operano nel recupero da residui da conceria.'
#    p=document.add_paragraph(T.format(Gruppo,df1.iloc[0][1]))
#    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#if len(df1)<2:
table = document.add_table(rows=len(df1)+2, cols=len(df1.columns)-8)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr1_cells = table.rows[1].cells
NNN=[COLS[4],COLS[5],COLS[7],COLS[9]]
hdr_cells[0].text = COLS[1]
hdr_cells[1].text = 'Sede Legale'
#hdr1_cells[1].text = COLS[4]
#hdr1_cells[2].text = COLS[5]
#hdr1_cells[3].text = COLS[7]
#hdr1_cells[4].text = COLS[9]
for i in range(1,5):
    my_cell = table.cell(1, i)
    my_paragraph = my_cell.paragraphs[0]
    run1 = my_paragraph.add_run(NNN[i-1])
    run1.font.color.rgb = WW

hdr_cells[5].text = COLS[10]
hdr_cells[6].text = COLS[12]
a = table.cell(0, 1)
b = table.cell(0, 4)
A = a.merge(b)
table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)

df1.index=range(0,len(df1))
df1=df1.fillna('-')
for i in range(0,len(df1)):
    table.cell(i+2, 0).text=str(df1['Denominazione societaria '].values[i])
    table.cell(i+2, 1).text=str(df1['Indirizzo'].values[i])
    table.cell(i+2, 2).text=str(df1['n.c.'].values[i])
    table.cell(i+2, 3).text=str(df1['città'].values[i])
    table.cell(i+2, 4).text=str(df1['nazione'].values[i])
    table.cell(i+2, 5).text=str(df1['Area di business'].values[i])
    table.cell(i+2, 6).text=str(df1['Capitale sociale'].values[i])

table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)

c=0
for i in table.columns:
    c+=1

for i in range(0,c):
    name='shading_elm_'+str(i)
    name=parse_xml(r'<w:shd {} w:fill="34657F"/>'.format(nsdecls('w')))
    table.rows[1].cells[i]._tc.get_or_add_tcPr().append(name)


document.add_heading('I Soci', level=2)

document.add_heading('Governance', level=2)
#%%%%%%%%%%%%%%%%%%%%%%%%
last_paragraph = document.paragraphs[-1]
last_paragraph.paragraph_format.space_before = Inches(1)
#%%%%%%%%%%%%%%%%%%%%%%%%
p=document.add_paragraph('Tutte le società sono guidate da un Consiglio di Amministrazione e da un Collegio sindacale')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


document.add_heading('Stakeholder', level=2)
p=document.add_paragraph('I principali stakeholder sono:')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph('Dipendenti', style='List Bullet')
document.add_paragraph('Associazioni di categoria', style='List Bullet')
document.add_paragraph('Fornitori', style='List Bullet')
document.add_paragraph('Azionisti e Finanziatori', style='List Bullet')
document.add_paragraph('Clienti', style='List Bullet')
document.add_paragraph('Comunità locale', style='List Bullet')
document.add_paragraph('Enti regolatori e authority', style='List Bullet')
document.add_paragraph('Università ed enti di ricerca', style='List Bullet')
p=document.add_paragraph('Gli interessi dei vari stakeholder sono di seguito sintetizzati:')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph('Dipendenti: benessere in ambiente lavorativo e proprio', style='List Bullet')
document.add_paragraph('Associazioni di categoria: sviluppo conoscenza e competenze', style='List Bullet')
document.add_paragraph('Clienti: Qualità, competitività del prodotto', style='List Bullet')
document.add_paragraph('Fornitori: Continuità operativa', style='List Bullet')
document.add_paragraph('Azionisti e finanziatori: redditività nel lungo periodo e sostenibilità del modello di business', style='List Bullet')
document.add_paragraph('Comunità locale: Sostegno economico e culturale', style='List Bullet')
document.add_paragraph('Enti regolatori: Compliance con normative', style='List Bullet')
document.add_paragraph('Università e centri di ricerca: Sostegno all’attività formativa e di ricerca', style='List Bullet')

document.add_heading('Piano strategico', level=2)
p=document.add_paragraph('Il piano strategico prevede il mantenimento della redditività nel lungo periodo mediante sviluppo di prodotti innovativi, sostegno dell’economia circolare, rispettando l’ambiente e sviluppando la resilienza.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_heading('Modello di business e posizionamento', level=2)
document.add_heading('Certificazioni', level=2)
#%%%%%%%%%%%%%%%%%%%%%%%%
last_paragraph = document.paragraphs[-1]
last_paragraph.paragraph_format.space_before = Inches(1)
#%%%%%%%%%%%%%%%%%%%%%%%%

if len(df1)<2:
    p=document.add_paragraph('La società ha conseguito le seguenti certificazioni:')
else:
    p=document.add_paragraph('Le società hanno conseguito le seguenti certificazioni:')

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_heading('Dati economico patrimoniali', level=2)
p=document.add_paragraph('La società mostra una redditività di rilievo e una struttura finanziaria solida.')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.add_heading('Prodotti', level=2)
document.add_heading('Marchi', level=2)
document.add_heading('Principali clienti', level=2)
document.add_heading('Distribuzione geografica dei ricavi', level=2)
document.add_heading('Sistemi di qualità', level=2)

document.add_page_break()
document.add_heading('Ubicazioni', level=2)

table = document.add_table(rows=len(df3)+2, cols=7)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr_cells1 = table.rows[1].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Ubicazione'
hdr_cells[6].text = 'Anno costruzione'
NNN=['società','indirizzo','n°','città','nazione']
for i in range(1,6):
    my_cell = table.cell(1, i)
    my_paragraph = my_cell.paragraphs[0]
    run1 = my_paragraph.add_run(NNN[i-1])
    run1.font.color.rgb = WW
#hdr_cells1[1].text = 'società'
#hdr_cells1[2].text = 'indirizzo'
#hdr_cells1[3].text = 'n°'
#hdr_cells1[4].text = 'città'
#hdr_cells1[5].text = 'nazione'

df3=df3.fillna('-')
for i in range(0,len(df3)):
    table.cell(i+2, 0).text=str(df3['ID'].values[i])
    table.cell(i+2, 1).text=str(df3['società'].values[i])
    table.cell(i+2, 2).text=str(df3['indirizzo'].values[i])
    table.cell(i+2, 3).text=str(df3['n°'].values[i])
    table.cell(i+2, 4).text=str(df3['città'].values[i])
    table.cell(i+2, 5).text=str(df3['nazione'].values[i])
    table.cell(i+2, 6).text=str(df3['Anno costruzione'].values[i])

a = table.cell(0, 1)
b = table.cell(0, 5)
A = a.merge(b)

table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)
c=0
for i in table.columns:
    c+=1

for i in range(0,c):
    name='shading_elm_'+str(i)
    name=parse_xml(r'<w:shd {} w:fill="34657F"/>'.format(nsdecls('w')))
    table.rows[1].cells[i]._tc.get_or_add_tcPr().append(name)

p=document.add_paragraph('')

table = document.add_table(rows=len(df3)+2, cols=8)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr_cells1 = table.rows[1].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'bene culturale (SI/NO)'
hdr_cells[2].text = 'Dati relativi all\'ubicazione'
hdr_cells[7].text = 'N. Dipendenti'
NNN=['soggetto proprietario','superficie di sviluppo totale (mq)','superficie di sviluppo specifica (mq)','destinazione d\'uso generale','destinazione d\'uso specifica']
for i in range(2,7):
    my_cell = table.cell(1, i)
    my_paragraph = my_cell.paragraphs[0]
    run1 = my_paragraph.add_run(NNN[i-2])
    run1.font.color.rgb = WW
#hdr_cells1[2].text = 'soggetto proprietario'
#hdr_cells1[3].text = 'superficie di sviluppo totale (mq)'
#hdr_cells1[4].text = 'superficie di sviluppo specifica (mq)'
#hdr_cells1[5].text = 'destinazione d\'uso generale'
#hdr_cells1[6].text = 'destinazione d\'uso specifica'

for i in range(0,len(df3)):
    table.cell(i+2, 0).text=str(df3['ID'].values[i])
    table.cell(i+2, 1).text=str(df3['bene culturale (SI/NO)'].values[i])
    table.cell(i+2, 2).text=str(df3['soggetto proprietario'].values[i])
    table.cell(i+2, 3).text=str(df3['superficie di sviluppo totale (mq) (1)'].values[i])
    table.cell(i+2, 4).text=str(df3['superficie di sviluppo specifica (mq) (1)'].values[i])
    table.cell(i+2, 5).text=str(df3['destinazione d\'uso generale (2)'].values[i])
    table.cell(i+2, 6).text=str(df3['destinazione d\'uso specifica (3)'].values[i])
    table.cell(i+2, 7).text=str(df3['N. Dipendenti'].values[i])

a = table.cell(0, 2)
b = table.cell(0, 6)
A = a.merge(b)
table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)
c=0
for i in table.columns:
    c+=1

for i in range(0,c):
    name='shading_elm_'+str(i)
    name=parse_xml(r'<w:shd {} w:fill="34657F"/>'.format(nsdecls('w')))
    table.rows[1].cells[i]._tc.get_or_add_tcPr().append(name)

p=document.add_paragraph('')

table = document.add_table(rows=len(df3)+2, cols=6)
table.style = 'Light Grid'
hdr_cells = table.rows[0].cells
hdr_cells1 = table.rows[1].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Somme attualmente assicurate o periziate'
NNN=['Fabbricato','Contenuto','Merci\n(max giacenza)','Totale','Polizza']
for i in range(1,6):
    my_cell = table.cell(1, i)
    my_paragraph = my_cell.paragraphs[0]
    run1 = my_paragraph.add_run(NNN[i-1])
    run1.font.color.rgb = WW
#hdr_cells1[1].text = 'Fabbricato'
#hdr_cells1[2].text = 'Contenuto'
#hdr_cells1[3].text = 'Merci\n(max giacenza)'
#hdr_cells1[4].text = 'Totale'
#hdr_cells1[5].text = 'Polizza'

for i in range(0,len(df3)):
    table.cell(i+2, 0).text=str(df3['ID'].values[i])
    table.cell(i+2, 1).text=str(df3['Fabbricato'].values[i])
    table.cell(i+2, 2).text=str(df3['Contenuto (5)'].values[i])
    table.cell(i+2, 3).text=str(df3['Merci\n(max giacenza)'].values[i])
    table.cell(i+2, 4).text=str(df3['Totale'].values[i])
    table.cell(i+2, 5).text=str(df3['Polizza'].values[i])


a = table.cell(0, 1)
b = table.cell(0, 5)
A = a.merge(b)
table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
table.rows[0].height=Inches(0.3)
c=0
for i in table.columns:
    c+=1

for i in range(0,c):
    name='shading_elm_'+str(i)
    name=parse_xml(r'<w:shd {} w:fill="34657F"/>'.format(nsdecls('w')))
    table.rows[1].cells[i]._tc.get_or_add_tcPr().append(name)

document.add_heading('Supply Chain', level=2)
p=document.add_paragraph('La supply chain è  caratterizzata da')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph('materie prime', style='List Bullet')
document.add_paragraph('energia', style='List Bullet')
document.add_paragraph('prodotti finiti utilizzati', style='List Bullet')





document.save('Rep_Client.docx')
